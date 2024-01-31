from docx import Document
from docx.shared import Pt
from io import BytesIO
from User import User
from datetime import datetime
import re
from SlackAPI import SlackAPI

def replace_user_ids_with_names(input_string, users: list[str, User]):
    # Define a regular expression pattern to match <@user_id>
    pattern = re.compile(r'<@(\w+)>')

    # Use a lambda function to perform the replacement
    result_string = pattern.sub(lambda match: str(users.get(match.group(1), match.group(0))), input_string)

    return result_string

def replace_links(input_string):
    pattern = re.compile(r'<([^|>]*)[^>]*>')
    result_string = pattern.sub(r'\1', input_string)
    return result_string

def print_message(message, users: list[str, User], indentation: bool, doc: Document, slack_api: SlackAPI):
    user_id = message.get('user')
    user = str(users.get(user_id, "Unknown"))
    user_paragraph = doc.add_paragraph(user)
    user_paragraph.runs[0].bold = True

    date_time = datetime.fromtimestamp(float(message["ts"]))
    date_time_pretty = date_time.strftime('%Y-%m-%d %H:%M:%S')

    user_paragraph.add_run(f', {date_time_pretty}')
    user_paragraph.paragraph_format.left_indent = Pt(20) if indentation else 0

    message_text = replace_user_ids_with_names(message["text"], users)
    message_text = replace_links(message_text)

    if len(message_text) > 0:
        message_paragraph = doc.add_paragraph(message_text)
        message_paragraph.paragraph_format.left_indent = Pt(20) if indentation else 0

    if message.get('files'):
        files_paragraph = doc.add_paragraph()
        files_paragraph.paragraph_format.left_indent = Pt(20) if indentation else 0

        files_run = files_paragraph.add_run(f'{len(message["files"])} files:\n')
        files_run.italic = True

        for file in message['files']:
            files_paragraph.add_run(f'{file["name"]} ({file["id"]})\n')

            file_data = slack_api.get_file(file) # this call will use cache because all files were already downloaded
            if file_data:
                if file["mimetype"].startswith("image"):
                    width = file["original_w"]
                    height = file["original_h"]
                    aspect_ratio = width / height

                    # Calculate the corresponding width to fit the maximum height
                    max_height = 400
                    max_width = 535

                    # Calculate dimensions to fit within the specified max_width and max_height
                    if width > max_width:
                        width = max_width
                        height = width / aspect_ratio

                    if height > max_height:
                        height = max_height
                        width = height * aspect_ratio

                    files_paragraph.add_run().add_picture(BytesIO(file_data), width=Pt(width), height=Pt(height))

def print_conversation_replies(data, users: list[str, User], doc: Document, slack_api: SlackAPI):

    replies_paragraph = doc.add_paragraph(f'{len(data)} replies:')
    replies_paragraph.runs[0].italic = True

    for message in data:
        if message.get('type', None) != 'message':
            continue
        print_message(message, users, indentation=True, doc=doc, slack_api=slack_api)

def print_conversation_history(data, users: list[str, User], channel_id: str, channel_name: str, slack_api: SlackAPI):
    doc = Document()

    section = doc.sections[0]

    # Set the page margins
    margin = 20
    section.left_margin = Pt(margin)
    section.right_margin = Pt(margin)
    section.top_margin = Pt(margin)
    section.bottom_margin = Pt(margin)

    replies_buffer = []

    idx = 0
    while idx < len(data):
        message = data[idx]

        print_message(message, users, indentation=False, doc=doc, slack_api=slack_api)

        if message.get('reply_count', 0) > 0:
            reply_count = message['reply_count']
            replies = data[idx + 1:idx+1+reply_count]
            idx += 1 + reply_count
            print_conversation_replies(replies, users, doc=doc, slack_api=slack_api)
        else:
            idx += 1
    
    doc.save(f'Output/{channel_name}/slack-export_{channel_name}.docx')